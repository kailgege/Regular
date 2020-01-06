'''
Created on 2019年11月18日

@author: 13476

print( u'中文'.encode('utf-8'))
'''
import re
from docx import Document

#f=Document('test/English.docx')
f=Document('test/zuowen.docx')
string=input("需要查找的字符串：")

re_f=re.compile(u'[.\w\W\u4e00-\u9fa5|\u3000-\u303f|\ufb00-\ufffd]*'+string+'[.\w\W\u4e00-\u9fa5||\u3000-\u303f|\ufb00-\ufffd]*',re.I) 
re_p=re.compile(r'\.|\?|\!|。|？|！|……')
for fparagraph in f.paragraphs:
    p=re.split(re_p,fparagraph.text)
    #print(p)
    for item in p:
        findall_f=re_f.findall(item)  
        for find_f in findall_f:
            print(find_f+'。')