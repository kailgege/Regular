'''
Created on 2019年11月14日

@author: 13476
'''
import re
from docx import Document

string1=input("需要替换的字符串：")
string2=input("替换的字符串：")

d=Document('test/AA.docx')
#d=Document('test/zuowen.docx')
for dparagraph in d.paragraphs:
    #dparagraph.text= dparagraph.text.replace('是','ABC')
    dparagraph.text=re.sub(string1,string2,dparagraph.text)
d.save('test/AA_out.docx')
#d.save('test/zuowen_out.docx')
print("替换成功")
